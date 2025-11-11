from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import shutil
from openpyxl import Workbook
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT, TA_CENTER
import json

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# تحديد مجلد المشروع
ROOT_DIR = Path(__file__).parent
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# إنشاء الراوتر مع بادئة /api
api_router = APIRouter(prefix="/api")

# الآن نعرف الروت بعد إنشاء الراوتر
@api_router.get("/test")
def test():
    return {"message": "API is working!"}

@app.get("/")
def root():
    return {"message": "API is working on the root!"}

app.include_router(api_router)



# ربط الراوتر بالتطبيق
app.include_router(api_router)
# Models
class CompanyInfo(BaseModel):
    name_ar: str = "شركة مثلث الأنظمة المميزة للمقاولات"
    name_en: str = "MUTHALLATH AL-ANZIMAH AL-MUMAYYIZAH CONTRACTING CO."
    description_ar: str = "تصميم وتصنيع وتوريد وتركيب مظلات الشد الإنشائي والخيام والسواتر"
    description_en: str = "Design, Manufacture, Supply & Installation of Structure Tension Awnings, Tents & Canopies"
    tax_number: str = "311104439400003"
    street: str = "شارع حائل"
    neighborhood: str = "حي البغدادية الغربية"
    country: str = "السعودية"
    city: str = "جدة"
    commercial_registration: str = "4030255240"
    building: str = "8376"
    postal_code: str = "22231"
    additional_number: str = "3842"
    email: str = "info@tsscoksa.com"
    phone1: str = "+966 50 061 2006"
    phone2: str = "055 538 9792"
    phone3: str = "+966 50 336 5527"
    logo_path: Optional[str] = None

class CustomerInfo(BaseModel):
    name: str
    tax_number: Optional[str] = None
    street: Optional[str] = None
    neighborhood: Optional[str] = None
    country: Optional[str] = "السعودية"
    city: Optional[str] = None
    commercial_registration: Optional[str] = None
    building: Optional[str] = None
    postal_code: Optional[str] = None
    additional_number: Optional[str] = None
    phone: Optional[str] = None

class QuoteItem(BaseModel):
    description: str
    quantity: float
    unit: str
    unit_price: float
    total_price: float

class QuoteCreate(BaseModel):
    customer: CustomerInfo
    project_description: str
    location: str
    items: List[QuoteItem]
    subtotal: float
    tax_amount: float
    total_amount: float
    notes: Optional[str] = None

class Quote(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    quote_number: str
    customer: CustomerInfo
    project_description: str
    location: str
    items: List[QuoteItem]
    subtotal: float
    tax_amount: float
    total_amount: float
    notes: Optional[str] = None
    created_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class QuoteUpdate(BaseModel):
    customer: Optional[CustomerInfo] = None
    project_description: Optional[str] = None
    location: Optional[str] = None
    items: Optional[List[QuoteItem]] = None
    subtotal: Optional[float] = None
    tax_amount: Optional[float] = None
    total_amount: Optional[float] = None
    notes: Optional[str] = None

# Utility functions
async def get_next_quote_number():
    """Generate next sequential quote number"""
    last_quote = await db.quotes.find_one(sort=[("created_date", -1)])
    if last_quote:
        try:
            last_number = int(last_quote.get("quote_number", "0"))
            return str(last_number + 1)
        except:
            count = await db.quotes.count_documents({})
            return str(count + 1)
    return "1"

# Routes
@api_router.get("/")
async def root():
    return {"message": "Quote Management System API"}

# Company routes
@api_router.get("/company", response_model=CompanyInfo)
async def get_company_info():
    company = await db.company.find_one({})
    if not company:
        # Create default company info
        default_company = CompanyInfo()
        await db.company.insert_one(default_company.dict())
        return default_company
    return CompanyInfo(**company)

@api_router.put("/company", response_model=CompanyInfo)
async def update_company_info(company: CompanyInfo):
    company_dict = company.dict()
    company_dict["updated_date"] = datetime.now(timezone.utc).isoformat()
    
    await db.company.delete_many({})  # Remove old company info
    await db.company.insert_one(company_dict)
    return company

@api_router.post("/company/logo")
async def upload_logo(file: UploadFile = File(...)):
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="File must be an image")
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    filename = f"logo_{uuid.uuid4()}{file_extension}"
    file_path = UPLOAD_DIR / filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Update company info with logo path
    await db.company.update_one(
        {},
        {"$set": {"logo_path": f"/api/uploads/{filename}"}},
        upsert=True
    )
    
    return {"logo_path": f"/api/uploads/{filename}"}

@api_router.get("/uploads/{filename}")
async def get_uploaded_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path)

# Quote routes
@api_router.post("/quotes", response_model=Quote)
async def create_quote(quote_data: QuoteCreate):
    quote_number = await get_next_quote_number()
    quote_dict = quote_data.dict()
    quote_dict["id"] = str(uuid.uuid4())
    quote_dict["quote_number"] = quote_number
    quote_dict["created_date"] = datetime.now(timezone.utc).isoformat()
    quote_dict["updated_date"] = datetime.now(timezone.utc).isoformat()
    
    quote_obj = Quote(**quote_dict)
    await db.quotes.insert_one(quote_obj.dict())
    return quote_obj

@api_router.get("/quotes", response_model=List[Quote])
async def get_quotes(skip: int = 0, limit: int = 100):
    quotes = await db.quotes.find().sort("created_date", -1).skip(skip).limit(limit).to_list(limit)
    return [Quote(**quote) for quote in quotes]

@api_router.get("/quotes/{quote_id}", response_model=Quote)
async def get_quote(quote_id: str):
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    return Quote(**quote)

@api_router.put("/quotes/{quote_id}", response_model=Quote)
async def update_quote(quote_id: str, quote_update: QuoteUpdate):
    existing_quote = await db.quotes.find_one({"id": quote_id})
    if not existing_quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    
    update_data = {k: v for k, v in quote_update.dict().items() if v is not None}
    update_data["updated_date"] = datetime.now(timezone.utc).isoformat()
    
    await db.quotes.update_one({"id": quote_id}, {"$set": update_data})
    
    updated_quote = await db.quotes.find_one({"id": quote_id})
    return Quote(**updated_quote)

@api_router.delete("/quotes/{quote_id}")
async def delete_quote(quote_id: str):
    result = await db.quotes.delete_one({"id": quote_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Quote not found")
    return {"message": "Quote deleted successfully"}

# Export routes
@api_router.get("/quotes/{quote_id}/export/excel")
async def export_quote_excel(quote_id: str):
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    
    quote_obj = Quote(**quote)
    company = await get_company_info()
    
    # Create Excel workbook
    wb = Workbook()
    ws = wb.active
    ws.title = f"Quote_{quote_obj.quote_number}"
    
    # Headers
    ws.append([f"Quote #{quote_obj.quote_number}"])
    ws.append([f"Company: {company.name_ar}"])
    ws.append([f"Customer: {quote_obj.customer.name}"])
    ws.append([f"Project: {quote_obj.project_description}"])
    ws.append([f"Location: {quote_obj.location}"])
    ws.append([])
    
    # Items table
    ws.append(["#", "Description", "Quantity", "Unit", "Unit Price", "Total"])
    for i, item in enumerate(quote_obj.items, 1):
        ws.append([i, item.description, item.quantity, item.unit, item.unit_price, item.total_price])
    
    ws.append([])
    ws.append(["", "", "", "", "Subtotal:", quote_obj.subtotal])
    ws.append(["", "", "", "", "Tax (15%):", quote_obj.tax_amount])
    ws.append(["", "", "", "", "Total:", quote_obj.total_amount])
    
    # Save to BytesIO
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    headers = {
        'Content-Disposition': f'attachment; filename="quote_{quote_obj.quote_number}.xlsx"'
    }
    
    return StreamingResponse(
        BytesIO(output.read()),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers
    )

@api_router.get("/quotes/{quote_id}/export/pdf")
async def export_quote_pdf(quote_id: str):
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    
    quote_obj = Quote(**quote)
    company = await get_company_info()
    
    # Create PDF using canvas to match exact preview layout
    buffer = BytesIO()
    
    # Use basic canvas approach for pixel-perfect matching
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors as pdf_colors
    from reportlab.pdfbase.pdfmetrics import registerFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import ImageReader
    import textwrap
    
    # Create canvas
    c = pdf_canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # Define measurements matching preview exactly
    margin_left = 20 * mm
    margin_right = 20 * mm 
    margin_top = 20 * mm
    margin_bottom = 20 * mm
    
    content_width = width - margin_left - margin_right
    y_position = height - margin_top
    
    def draw_text_right_aligned(canvas, text, x, y, font_name="Helvetica", font_size=10):
        """Draw right-aligned text with UTF-8 support"""
        canvas.setFont(font_name, font_size)
        text_width = canvas.stringWidth(str(text), font_name, font_size)
        canvas.drawString(x - text_width, y, str(text))
        return text_width
    
    def draw_text_center_aligned(canvas, text, x, y, font_name="Helvetica", font_size=10):
        """Draw center-aligned text"""
        canvas.setFont(font_name, font_size)
        text_width = canvas.stringWidth(str(text), font_name, font_size)
        canvas.drawString(x - text_width/2, y, str(text))
    
    def draw_bordered_box(canvas, x, y, width, height, fill_color=None):
        """Draw a bordered box"""
        if fill_color:
            canvas.setFillColor(fill_color)
            canvas.rect(x, y - height, width, height, fill=1, stroke=1)
            canvas.setFillColor(pdf_colors.black)
        else:
            canvas.rect(x, y - height, width, height, fill=0, stroke=1)
    
    # === PAGE 1: HEADER SECTION (exactly like preview) ===
    
    # Header layout: Logo | Company Info | Quote Badge
    header_height = 80
    
    # Logo area (left side)
    logo_x = margin_left
    logo_width = 60
    if company.logo_path:
        # Would draw logo here if path exists
        pass
    
    # Company info (center)
    company_center_x = margin_left + content_width / 2
    
    # Company name (Arabic) - large bold
    c.setFont("Helvetica-Bold", 18)
    draw_text_center_aligned(c, company.name_ar or "شركة مثلث الأنظمة المميزة للمقاولات", 
                           company_center_x, y_position - 10)
    
    # Company description (Arabic)
    c.setFont("Helvetica", 11)
    draw_text_center_aligned(c, company.description_ar, company_center_x, y_position - 30)
    
    # Company name (English)
    c.setFont("Helvetica", 9)
    draw_text_center_aligned(c, company.name_en, company_center_x, y_position - 45)
    
    # Quote badge (right side)
    quote_x = margin_left + content_width
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(pdf_colors.blue)
    draw_text_right_aligned(c, f"عرض سعر رقم {quote_obj.quote_number}", 
                          quote_x, y_position - 10)
    
    c.setFillColor(pdf_colors.black)
    c.setFont("Helvetica", 10)
    draw_text_right_aligned(c, quote_obj.created_date.strftime("%B %d, %Y"), 
                          quote_x, y_position - 30)
    
    y_position -= 100
    
    # === COMPANY & CUSTOMER INFO TABLE (exactly like preview) ===
    
    # Draw table header with blue background
    table_y = y_position
    col_width = content_width / 2
    row_height = 25
    
    # Header row with blue background
    draw_bordered_box(c, margin_left, table_y, col_width, row_height, pdf_colors.lightblue)
    draw_bordered_box(c, margin_left + col_width, table_y, col_width, row_height, pdf_colors.lightblue)
    
    c.setFont("Helvetica-Bold", 12)
    draw_text_center_aligned(c, "Seller / المورد", margin_left + col_width/2, table_y - 15)
    draw_text_center_aligned(c, "Customer / العميل", margin_left + col_width*1.5, table_y - 15)
    
    # Data rows
    company_data = [
        (f"الشركة: {company.name_ar}", f"العميل: {quote_obj.customer.name}"),
        (f"الرقم الضريبي: {company.tax_number}", f"الرقم الضريبي: {quote_obj.customer.tax_number or 'غير محدد'}"),
        (f"الشارع: {company.street}", f"الشارع: {quote_obj.customer.street or 'غير محدد'}"),
        (f"الحي: {company.neighborhood}", f"الحي: {quote_obj.customer.neighborhood or 'غير محدد'}"),
        (f"المدينة: {company.city}", f"المدينة: {quote_obj.customer.city or 'غير محدد'}"),
        (f"الدولة: {company.country}", f"الدولة: {quote_obj.customer.country or 'السعودية'}"),
        (f"السجل التجاري: {company.commercial_registration}", f"السجل التجاري: {quote_obj.customer.commercial_registration or 'غير محدد'}"),
        (f"المبنى: {company.building}", f"المبنى: {quote_obj.customer.building or 'غير محدد'}"),
        (f"الرمز البريدي: {company.postal_code}", f"الرمز البريدي: {quote_obj.customer.postal_code or 'غير محدد'}"),
        (f"الرقم الإضافي: {company.additional_number}", f"الرقم الإضافي: {quote_obj.customer.additional_number or 'غير محدد'}"),
        ("", f"رقم الهاتف: {quote_obj.customer.phone or 'غير محدد'}")
    ]
    
    row_y = table_y - row_height
    c.setFont("Helvetica", 9)
    
    for seller_text, customer_text in company_data:
        # Draw row borders
        draw_bordered_box(c, margin_left, row_y, col_width, 18)
        draw_bordered_box(c, margin_left + col_width, row_y, col_width, 18)
        
        # Draw text (right-aligned)
        if seller_text:
            draw_text_right_aligned(c, seller_text, margin_left + col_width - 5, row_y - 12)
        if customer_text:
            draw_text_right_aligned(c, customer_text, margin_left + col_width*2 - 5, row_y - 12)
        
        row_y -= 18
    
    y_position = row_y - 20
    
    # === PROJECT DETAILS SECTION ===
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(pdf_colors.purple)
    draw_text_right_aligned(c, "تفاصيل المشروع / Project details", 
                          margin_left + content_width, y_position)
    
    y_position -= 30
    c.setFillColor(pdf_colors.black)
    
    # Project details box with gray background
    project_box_height = 50
    draw_bordered_box(c, margin_left, y_position, content_width, project_box_height, 
                     pdf_colors.Color(0.95, 0.95, 0.95))  # Light gray
    
    c.setFont("Helvetica-Bold", 10)
    draw_text_right_aligned(c, "وصف المشروع:", margin_left + 80, y_position - 15)
    c.setFont("Helvetica", 10)
    
    # Wrap project description if too long
    desc_lines = textwrap.wrap(quote_obj.project_description, width=80)
    for i, line in enumerate(desc_lines[:3]):  # Max 3 lines
        draw_text_right_aligned(c, line, margin_left + content_width - 5, y_position - 15 - i*12)
    
    c.setFont("Helvetica-Bold", 10)
    draw_text_right_aligned(c, "الموقع:", margin_left + 50, y_position - 40)
    c.setFont("Helvetica", 10)
    draw_text_right_aligned(c, quote_obj.location or "غير محدد", 
                          margin_left + content_width - 5, y_position - 40)
    
    y_position -= 80
    
    # === ITEMS TABLE SECTION ===
    c.setFont("Helvetica-Bold", 14) 
    c.setFillColor(pdf_colors.purple)
    draw_text_right_aligned(c, "بنود عرض السعر / Price table", 
                          margin_left + content_width, y_position)
    
    y_position -= 30
    c.setFillColor(pdf_colors.black)
    
    # Table headers
    col_widths = [25, 90, 25, 30, 35, 35]  # in mm
    col_positions = []
    current_x = margin_left
    for width in col_widths:
        col_positions.append(current_x)
        current_x += width * mm
    
    headers = ["الرقم التسلسلي", "الوصف", "الكمية", "الوحدة", "سعر الوحدة", "السعر الإجمالي"]
    
    # Draw header row with gray background
    table_width = sum(col_widths) * mm
    header_row_height = 20
    draw_bordered_box(c, margin_left, y_position, table_width, header_row_height, 
                     pdf_colors.grey)
    
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(pdf_colors.white)
    for i, header in enumerate(headers):
        header_x = col_positions[i] + col_widths[i] * mm / 2
        draw_text_center_aligned(c, header, header_x, y_position - 13)
    
    c.setFillColor(pdf_colors.black)
    y_position -= header_row_height
    
    # Data rows
    items_per_page = 15  # Limit to prevent page overflow
    row_height = 20
    
    for i, item in enumerate(quote_obj.items[:items_per_page], 1):
        # Check if we need a new page
        if y_position - row_height < margin_bottom + 100:  # Leave space for totals
            c.showPage()  # New page
            y_position = height - margin_top
            
            # Redraw headers on new page
            draw_bordered_box(c, margin_left, y_position, table_width, header_row_height, 
                             pdf_colors.grey)
            c.setFont("Helvetica-Bold", 10)
            c.setFillColor(pdf_colors.white)
            for j, header in enumerate(headers):
                header_x = col_positions[j] + col_widths[j] * mm / 2
                draw_text_center_aligned(c, header, header_x, y_position - 13)
            c.setFillColor(pdf_colors.black)
            y_position -= header_row_height
        
        # Draw row border
        for j in range(len(col_widths)):
            draw_bordered_box(c, col_positions[j], y_position, col_widths[j] * mm, row_height)
        
        # Draw data
        c.setFont("Helvetica", 9)
        
        # Serial number (center)
        draw_text_center_aligned(c, str(i), col_positions[0] + col_widths[0] * mm / 2, y_position - 13)
        
        # Description (right-aligned, truncated if needed)
        desc = item.description[:35] + "..." if len(item.description) > 35 else item.description
        draw_text_right_aligned(c, desc, col_positions[1] + col_widths[1] * mm - 3, y_position - 13)
        
        # Quantity (center)
        draw_text_center_aligned(c, f"{item.quantity:g}", col_positions[2] + col_widths[2] * mm / 2, y_position - 13)
        
        # Unit (center)
        draw_text_center_aligned(c, item.unit, col_positions[3] + col_widths[3] * mm / 2, y_position - 13)
        
        # Unit price (center)
        draw_text_center_aligned(c, f"{item.unit_price:,.2f}", col_positions[4] + col_widths[4] * mm / 2, y_position - 13)
        
        # Total price (center)
        draw_text_center_aligned(c, f"{item.total_price:,.2f}", col_positions[5] + col_widths[5] * mm / 2, y_position - 13)
        
        y_position -= row_height
    
    # Handle remaining items if any (continue on next pages)
    if len(quote_obj.items) > items_per_page:
        for chunk_start in range(items_per_page, len(quote_obj.items), items_per_page):
            c.showPage()
            c.translate(0, 10*mm)  # 10mm space between pages
            y_position = height - margin_top - 10*mm
            
            chunk_end = min(chunk_start + items_per_page, len(quote_obj.items))
            
            # Draw headers
            draw_bordered_box(c, margin_left, y_position, table_width, header_row_height, 
                             pdf_colors.grey)
            c.setFont("Helvetica-Bold", 10)
            c.setFillColor(pdf_colors.white)
            for j, header in enumerate(headers):
                header_x = col_positions[j] + col_widths[j] * mm / 2
                draw_text_center_aligned(c, header, header_x, y_position - 13)
            c.setFillColor(pdf_colors.black)
            y_position -= header_row_height
            
            # Draw items for this chunk
            for i in range(chunk_start, chunk_end):
                item = quote_obj.items[i]
                
                # Draw row border
                for j in range(len(col_widths)):
                    draw_bordered_box(c, col_positions[j], y_position, col_widths[j] * mm, row_height)
                
                # Draw data
                c.setFont("Helvetica", 9)
                draw_text_center_aligned(c, str(i + 1), col_positions[0] + col_widths[0] * mm / 2, y_position - 13)
                
                desc = item.description[:35] + "..." if len(item.description) > 35 else item.description
                draw_text_right_aligned(c, desc, col_positions[1] + col_widths[1] * mm - 3, y_position - 13)
                
                draw_text_center_aligned(c, f"{item.quantity:g}", col_positions[2] + col_widths[2] * mm / 2, y_position - 13)
                draw_text_center_aligned(c, item.unit, col_positions[3] + col_widths[3] * mm / 2, y_position - 13)
                draw_text_center_aligned(c, f"{item.unit_price:,.2f}", col_positions[4] + col_widths[4] * mm / 2, y_position - 13)
                draw_text_center_aligned(c, f"{item.total_price:,.2f}", col_positions[5] + col_widths[5] * mm / 2, y_position - 13)
                
                y_position -= row_height
    
    y_position -= 30
    
    # === TOTALS SECTION (right-aligned like preview) ===
    totals_width = 120 * mm
    totals_x = margin_left + content_width - totals_width
    
    totals_data = [
        ("المجموع الفرعي:", f"{quote_obj.subtotal:,.2f} ريال"),
        ("ضريبة القيمة المضافة (15%):", f"{quote_obj.tax_amount:,.2f} ريال"),
        ("المبلغ الإجمالي:", f"{quote_obj.total_amount:,.2f} ريال")
    ]
    
    totals_row_height = 20
    
    for i, (label, amount) in enumerate(totals_data):
        # Draw borders
        draw_bordered_box(c, totals_x, y_position, 80*mm, totals_row_height)
        draw_bordered_box(c, totals_x + 80*mm, y_position, 40*mm, totals_row_height)
        
        # Highlight total row
        if i == 2:  # Total row
            draw_bordered_box(c, totals_x, y_position, 120*mm, totals_row_height, 
                             pdf_colors.Color(0.9, 0.9, 0.9))
            c.setFont("Helvetica-Bold", 12)
            c.setFillColor(pdf_colors.green)
        else:
            c.setFont("Helvetica", 11)
            c.setFillColor(pdf_colors.black)
        
        # Draw text
        draw_text_right_aligned(c, label, totals_x + 75*mm, y_position - 13)
        draw_text_right_aligned(c, amount, totals_x + 115*mm, y_position - 13)
        
        y_position -= totals_row_height
    
    c.setFillColor(pdf_colors.black)
    y_position -= 30
    
    # === NOTES SECTION ===
    if quote_obj.notes:
        c.setFont("Helvetica-Bold", 12)
        draw_text_right_aligned(c, "ملاحظات", margin_left + content_width, y_position)
        y_position -= 25
        
        # Yellow notes box
        notes_height = 60
        draw_bordered_box(c, margin_left, y_position, content_width, notes_height, 
                         pdf_colors.Color(1, 0.95, 0.8))  # Light yellow
        
        c.setFont("Helvetica", 10)
        # Wrap notes text
        notes_lines = textwrap.wrap(quote_obj.notes, width=100)
        for i, line in enumerate(notes_lines[:4]):  # Max 4 lines
            draw_text_right_aligned(c, line, margin_left + content_width - 10, y_position - 15 - i*12)
        
        y_position -= 80
    
    # === SIGNATURE SECTION ===
    y_position -= 20
    
    c.setFont("Helvetica-Bold", 12)
    draw_text_center_aligned(c, "التوقيع والاعتماد", margin_left + content_width/2, y_position)
    y_position -= 30
    
    # Signature table
    sig_col_width = content_width / 2
    sig_row_height = 25
    
    # Header
    draw_bordered_box(c, margin_left, y_position, sig_col_width, sig_row_height, 
                     pdf_colors.Color(0.9, 0.9, 0.9))
    draw_bordered_box(c, margin_left + sig_col_width, y_position, sig_col_width, sig_row_height, 
                     pdf_colors.Color(0.9, 0.9, 0.9))
    
    c.setFont("Helvetica-Bold", 11)
    draw_text_center_aligned(c, "التوقيع والختم", margin_left + sig_col_width/2, y_position - 15)
    draw_text_center_aligned(c, "تاريخ الموافقة", margin_left + sig_col_width*1.5, y_position - 15)
    
    # Empty signature rows
    for i in range(3):
        y_position -= sig_row_height
        draw_bordered_box(c, margin_left, y_position, sig_col_width, sig_row_height)
        draw_bordered_box(c, margin_left + sig_col_width, y_position, sig_col_width, sig_row_height)
    
    y_position -= 40
    
    # === FOOTER CONTACT INFO ===
    c.setFont("Helvetica-Bold", 10)
    draw_text_center_aligned(c, "معلومات الاتصال", margin_left + content_width/2, y_position)
    
    y_position -= 20
    c.setFont("Helvetica", 9)
    c.setFillColor(pdf_colors.grey)
    
    contact_lines = [
        f"البريد الإلكتروني: {company.email}",
        f"{company.neighborhood}, {company.city}",
        f"جوال: {company.phone1} | جوال آخر: {company.phone2 or ''} | جوال إضافي: {company.phone3 or ''}"
    ]
    
    for line in contact_lines:
        draw_text_center_aligned(c, line, margin_left + content_width/2, y_position)
        y_position -= 12
    
    # Save PDF
    c.save()
    buffer.seek(0)
    
    import time
    timestamp = int(time.time())
    headers = {
        'Content-Disposition': f'attachment; filename="quote_{quote_obj.quote_number}_v{timestamp}.pdf"',
        'Cache-Control': 'no-cache, no-store, must-revalidate',
        'Pragma': 'no-cache',
        'Expires': '0'
    }
    
    return StreamingResponse(
        BytesIO(buffer.read()),
        media_type="application/pdf",
        headers=headers
    )

# Word export route - matches preview layout
@api_router.get("/quotes/{quote_id}/export/word")
async def export_quote_word(quote_id: str):
    quote = await db.quotes.find_one({"id": quote_id})
    if not quote:
        raise HTTPException(status_code=404, detail="Quote not found")
    
    quote_obj = Quote(**quote)
    company = await get_company_info()
    
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    # Create Word document with RTL support
    doc = Document()
    
    # Set page layout to A4
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    
    # === HEADER SECTION - matches preview ===
    
    # Company header table (logo area, company info, quote info)
    header_table = doc.add_table(rows=3, cols=3)
    header_table.columns[0].width = Inches(1.5)  # Logo area
    header_table.columns[1].width = Inches(4.5)  # Company info
    header_table.columns[2].width = Inches(1.5)  # Quote info
    
    # Company name and info (center column)
    header_table.cell(0, 1).text = f'شركة {company.name_ar}'
    header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(16)
    header_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    header_table.cell(1, 1).text = company.description_ar
    header_table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.cell(1, 1).paragraphs[0].runs[0].font.size = Pt(10)
    
    header_table.cell(2, 1).text = company.name_en
    header_table.cell(2, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_table.cell(2, 1).paragraphs[0].runs[0].font.size = Pt(9)
    
    # Quote number and date (right column)
    header_table.cell(0, 2).text = f'عرض سعر رقم {quote_obj.quote_number}'
    header_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_table.cell(0, 2).paragraphs[0].runs[0].font.size = Pt(12)
    header_table.cell(0, 2).paragraphs[0].runs[0].font.bold = True
    
    header_table.cell(1, 2).text = f'التاريخ: {quote_obj.created_date.strftime("%B %d, %Y")}'
    header_table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_table.cell(1, 2).paragraphs[0].runs[0].font.size = Pt(10)
    
    # Remove borders from header table
    for row in header_table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
            )
    
    doc.add_paragraph()
    
    # === COMPANY AND CUSTOMER INFO TABLE - exactly like preview ===
    
    info_table = doc.add_table(rows=13, cols=2)
    info_table.style = 'Table Grid'
    
    # Header row
    info_table.cell(0, 0).text = 'Seller / المورد'
    info_table.cell(0, 1).text = 'Customer / العميل'
    
    # Style header row
    for i in range(2):
        cell = info_table.cell(0, i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        # Add blue background
        shading = parse_xml(r'<w:shd {} w:fill="D6EAF8"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
    
    # Data rows - matching preview layout exactly
    company_data = [
        (f'الشركة: {company.name_ar}', f'العميل: {quote_obj.customer.name}'),
        (f'الرقم الضريبي: {company.tax_number}', f'الرقم الضريبي: {quote_obj.customer.tax_number or "غير محدد"}'),
        (f'الشارع: {company.street}', f'الشارع: {quote_obj.customer.street or "غير محدد"}'),
        (f'الحي: {company.neighborhood}', f'الحي: {quote_obj.customer.neighborhood or "غير محدد"}'),
        (f'المدينة: {company.city}', f'المدينة: {quote_obj.customer.city or "غير محدد"}'),
        (f'الدولة: {company.country}', f'الدولة: {quote_obj.customer.country or "غير محدد"}'),
        (f'السجل التجاري: {company.commercial_registration}', f'السجل التجاري: {quote_obj.customer.commercial_registration or "غير محدد"}'),
        (f'المبنى: {company.building}', f'المبنى: {quote_obj.customer.building or "غير محدد"}'),
        (f'الرمز البريدي: {company.postal_code}', f'الرمز البريدي: {quote_obj.customer.postal_code or "غير محدد"}'),
        (f'الرقم الإضافي: {company.additional_number}', f'الرقم الإضافي: {quote_obj.customer.additional_number or "غير محدد"}'),
        ('', f'رقم الهاتف: {quote_obj.customer.phone or "غير محدد"}'),
        ('', '')
    ]
    
    for i, (company_text, customer_text) in enumerate(company_data, 1):
        info_table.cell(i, 0).text = company_text
        info_table.cell(i, 1).text = customer_text
        
        # Right-align text
        info_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        info_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Font size
        info_table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(9)
        info_table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # === PROJECT DETAILS ===
    project_heading = doc.add_heading('Project details / تفاصيل المشروع', level=2)
    project_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    project_heading.runs[0].font.color.rgb = RGBColor(0, 100, 200)  # Blue color
    
    # Project info table
    project_table = doc.add_table(rows=2, cols=2)
    project_table.columns[0].width = Inches(1.5)
    project_table.columns[1].width = Inches(5)
    
    project_table.cell(0, 0).text = 'وصف المشروع:'
    project_table.cell(0, 1).text = quote_obj.project_description
    project_table.cell(1, 0).text = 'الموقع:'
    project_table.cell(1, 1).text = quote_obj.location
    
    # Style project table
    for row in project_table.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Remove borders from project table
    for row in project_table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                parse_xml(r'<w:tcBorders %s><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>' % nsdecls('w'))
            )
    
    doc.add_paragraph()
    
    # === ITEMS TABLE - exactly like preview ===
    items_heading = doc.add_heading('Price table / جدول الأسعار', level=2)
    items_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    items_heading.runs[0].font.color.rgb = RGBColor(0, 100, 200)  # Blue color
    
    # Calculate if we need page breaks (approximately 20 items per page)
    items_per_page = 20
    total_items = len(quote_obj.items)
    
    if total_items > items_per_page:
        # Split into chunks for page breaks
        for chunk_start in range(0, total_items, items_per_page):
            chunk_end = min(chunk_start + items_per_page, total_items)
            
            # Create table for this chunk
            items_table = doc.add_table(rows=1, cols=6)
            items_table.style = 'Table Grid'
            
            # Header row
            headers = items_table.rows[0].cells
            headers[0].text = 'الرقم التسلسلي'
            headers[1].text = 'الوصف'
            headers[2].text = 'الكمية'
            headers[3].text = 'الوحدة'
            headers[4].text = 'سعر الوحدة'
            headers[5].text = 'السعر الإجمالي'
            
            # Style header
            for cell in headers:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                # Gray background for header
                shading = parse_xml(r'<w:shd {} w:fill="808080"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White text
            
            # Add items for this chunk
            for i in range(chunk_start, chunk_end):
                item = quote_obj.items[i]
                row_cells = items_table.add_row().cells
                row_cells[0].text = str(i + 1)
                row_cells[1].text = item.description
                row_cells[2].text = f"{item.quantity:g}"
                row_cells[3].text = item.unit
                row_cells[4].text = f"{item.unit_price:,.2f}"
                row_cells[5].text = f"{item.total_price:,.2f}"
                
                # Style data cells
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for j in range(2, 6):
                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Font size for data
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Add page break if not the last chunk
            if chunk_end < total_items:
                doc.add_page_break()
    else:
        # Single table for all items
        items_table = doc.add_table(rows=1, cols=6)
        items_table.style = 'Table Grid'
        
        # Header row
        headers = items_table.rows[0].cells
        headers[0].text = 'الرقم التسلسلي'
        headers[1].text = 'الوصف'
        headers[2].text = 'الكمية'
        headers[3].text = 'الوحدة'
        headers[4].text = 'سعر الوحدة'
        headers[5].text = 'السعر الإجمالي'
        
        # Style header
        for cell in headers:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            shading = parse_xml(r'<w:shd {} w:fill="808080"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Add all items
        for i, item in enumerate(quote_obj.items, 1):
            row_cells = items_table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = item.description
            row_cells[2].text = f"{item.quantity:g}"
            row_cells[3].text = item.unit
            row_cells[4].text = f"{item.unit_price:,.2f}"
            row_cells[5].text = f"{item.total_price:,.2f}"
            
            # Style data cells
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for j in range(2, 6):
                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for cell in row_cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # === TOTALS SECTION - matches preview ===
    
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Table Grid'
    totals_table.columns[0].width = Inches(3)
    totals_table.columns[1].width = Inches(2.5)
    
    totals_data = [
        ('المجموع الفرعي:', f'{quote_obj.subtotal:,.2f} ريال'),
        ('ضريبة القيمة المضافة (15%):', f'{quote_obj.tax_amount:,.2f} ريال'),
        ('المبلغ الإجمالي:', f'{quote_obj.total_amount:,.2f} ريال')
    ]
    
    for i, (desc, amount) in enumerate(totals_data):
        totals_table.cell(i, 0).text = desc
        totals_table.cell(i, 1).text = amount
        
        # Right align
        totals_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Font styling
        totals_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        totals_table.cell(i, 1).paragraphs[0].runs[0].font.bold = True
        totals_table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(11)
        totals_table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(11)
        
        # Highlight total row
        if i == 2:  # Total row
            totals_table.cell(i, 0).paragraphs[0].runs[0].font.size = Pt(13)
            totals_table.cell(i, 1).paragraphs[0].runs[0].font.size = Pt(13)
            # Gray background for total
            for j in range(2):
                shading = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
                totals_table.cell(i, j)._element.get_or_add_tcPr().append(shading)
    
    # Right-align totals table
    totals_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # === SIGNATURE SECTION - matches preview ===
    
    signature_heading = doc.add_heading('التوقيع والختم', level=2)
    signature_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    signature_table = doc.add_table(rows=4, cols=2)
    signature_table.style = 'Table Grid'
    
    # Header row
    signature_table.cell(0, 0).text = 'التوقيع والختم'
    signature_table.cell(0, 1).text = 'تاريخ الموافقة'
    
    # Style header
    for i in range(2):
        cell = signature_table.cell(0, i)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        shading = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading)
    
    # Empty rows for signatures
    for i in range(1, 4):
        for j in range(2):
            signature_table.cell(i, j).text = ""
    
    # Notes section if exists
    if quote_obj.notes:
        doc.add_paragraph()
        notes_heading = doc.add_heading('ملاحظات', level=2)
        notes_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        notes_para = doc.add_paragraph(quote_obj.notes)
        notes_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Contact footer
    doc.add_paragraph()
    footer_heading = doc.add_heading('معلومات الاتصال', level=3)
    footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_text = f"""البريد الإلكتروني: {company.email}
{company.neighborhood}, {company.city}
جوال: {company.phone1} | جوال آخر: {company.phone2 or ''} | جوال إضافي: {company.phone3 or ''}"""
    
    contact_para = doc.add_paragraph(contact_text)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(9)
    
    # Save document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    timestamp = int(time.time())
    headers = {
        'Content-Disposition': f'attachment; filename="quote_{quote_obj.quote_number}_v{timestamp}.docx"',
        'Cache-Control': 'no-cache, no-store, must-revalidate',
        'Pragma': 'no-cache',
        'Expires': '0'
    }
    
    return StreamingResponse(
        BytesIO(buffer.read()),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()